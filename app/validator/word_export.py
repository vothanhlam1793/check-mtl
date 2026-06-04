import os
import shutil
import re
from datetime import datetime
from typing import List, Optional
from collections import defaultdict

from docx import Document
from docx.shared import Pt

from app.schemas import ErrorItem
from app.validator.engine import run_validation
from config import LLM_SUMMARIZE_THRESHOLD

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "..", "template-bc-tham-dinh.docx")


def _extract_project_info(file_name: str) -> dict:
    info = {"ten_du_an": "", "khu_vuc": "", "nhom_du_an": ""}
    fn_upper = file_name.upper()

    pt_match = re.search(r"PHUOC\s*THIEN\s*(\d+[A-Z]*)", fn_upper)
    tt_match = re.search(r"TAI\s*TIEN", fn_upper)
    ln_match = re.search(r"LA\s*NGA|KDT\s*LA\s*NGA", fn_upper)
    tro_match = re.search(r"TROPICANA", fn_upper)
    gm_match = re.search(r"GRAND\s*MERCURE", fn_upper)
    mor_match = re.search(r"MORITO|HO\s*TRAM", fn_upper)
    hl_match = re.search(r"HOANG\s*LONG", fn_upper)
    hab_match = re.search(r"HABANA", fn_upper)
    rmn_match = re.search(r"RMN", fn_upper)

    if pt_match:
        info["ten_du_an"] = f"Phước Thiền {pt_match.group(1)}"
    elif tt_match:
        info["ten_du_an"] = "Tài Tiến"
    elif ln_match:
        info["ten_du_an"] = "La Nga"
    elif tro_match:
        info["ten_du_an"] = "Tropicana"
    elif gm_match:
        info["ten_du_an"] = "Grand Mercure"
    elif mor_match:
        info["ten_du_an"] = "Morito Hồ Tràm"
    elif hl_match:
        info["ten_du_an"] = "Hoàng Long"
    elif hab_match:
        info["ten_du_an"] = "Habana Island"
    elif rmn_match:
        info["ten_du_an"] = "RMN"
    else:
        info["ten_du_an"] = file_name.rsplit(".", 1)[0][:50]

    area_map = {
        "PHUOC THIEN": "Đồng Nai 2", "TAI TIEN": "Đồng Nai 2", "LA NGA": "Đồng Nai 2",
        "TROPICANA": "Bà Rịa - Vũng Tàu", "GRAND MERCURE": "Bà Rịa - Vũng Tàu",
        "MORITO": "Bà Rịa - Vũng Tàu", "HO TRAM": "Bà Rịa - Vũng Tàu",
        "HOANG LONG": "Bà Rịa - Vũng Tàu", "HABANA": "Bà Rịa - Vũng Tàu",
        "RMN": "Bà Rịa - Vũng Tàu",
    }
    info["khu_vuc"] = ""
    for key, val in area_map.items():
        if key in fn_upper:
            info["khu_vuc"] = val
            break

    group_match = re.search(r'DN(\d+)', fn_upper)
    if group_match:
        info["nhom_du_an"] = group_match.group(1)
    return info


def _set_cell_text(cell, text: str, bold: bool = False, font_name: str = "Tahoma", font_size: float = 10):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.clear()
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold


def _add_row(table, cells_data: List[dict]):
    row = table.add_row()
    for ci, data in enumerate(cells_data):
        if ci < len(row.cells):
            _set_cell_text(row.cells[ci], data.get("text", ""), bold=data.get("bold", False))


def _remove_rows_from(table, start_idx: int):
    while len(table.rows) > start_idx:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)


def _build_cover_rows(cover_errors: List[ErrorItem]) -> List[List[dict]]:
    """Build table rows for Cover errors (Section A) — with LLM summary."""
    rows = []
    crits = [e for e in cover_errors if e.severity == "CRITICAL"]
    warns = [e for e in cover_errors if e.severity == "WARNING"]

    if not crits and not warns:
        rows.append(_make_row(content="Thông tin dự án trên Cover", opinion="Đầy đủ, không có lỗi."))
        return rows

    # Build structured error list for LLM
    error_list = []
    for e in crits:
        error_list.append({"field": e.field, "reason": e.reason, "severity": "CRITICAL"})
    for e in warns:
        error_list.append({"field": e.field, "reason": e.reason, "severity": "WARNING"})

    # Try LLM summary
    llm_opinion = ""
    try:
        from app.renderer import summarize_cover_errors
        llm_opinion = summarize_cover_errors(error_list)
    except Exception:
        pass

    if llm_opinion:
        rows.append(_make_row(
            content="Đánh giá Cover",
            opinion=llm_opinion,
        ))
    else:
        # Fallback: list individually
        for e in crits:
            rows.append(_make_row(content=f"{e.field}", opinion=f"[NGHIÊM TRỌNG] {e.reason}"))
        for e in warns:
            rows.append(_make_row(content=f"{e.field}", opinion=f"[Nhắc nhở] {e.reason}"))

    return rows


def _make_row(wbs_text: str = "", content: str = "", opinion: str = "", bold: bool = False):
    """Helper to make a 7-cell row."""
    return [
        {"text": wbs_text, "bold": bold},
        {"text": ""},
        {"text": ""},
        {"text": content, "bold": bold},
        {"text": opinion},
        {"text": ""},
        {"text": ""},
    ]


def _build_wbs_sections(file_path: str, errors: List[ErrorItem]) -> List[List[dict]]:
    """Build Section B rows with full WBS hierarchy from Excel data."""
    from app.validator.file_check import check_and_open_file

    data_errors = [e for e in errors if e.col != "Cover" and e.severity in ("CRITICAL", "ERROR", "WARNING")]
    sev_labels = {"CRITICAL": "[NGHIÊM TRỌNG]", "ERROR": "[Lỗi]", "WARNING": "[Nhắc nhở]"}

    try:
        wb, ds_name, cs_name, layout = check_and_open_file(file_path)
        ws = wb[ds_name]
    except Exception:
        # Fallback: error-only listing
        return _fallback_data_rows(data_errors, sev_labels)

    # Read all WBS entries
    entries = []  # (row, wbs, task)
    for r in range(layout.header_row + 1, ws.max_row + 1):
        wbs_v = ws.cell(row=r, column=layout.wbs_col + 1).value
        task_v = ws.cell(row=r, column=layout.task_col + 1).value
        if wbs_v is not None and str(wbs_v).strip():
            entries.append((r, str(wbs_v).strip(), str(task_v).strip() if task_v else ""))
    wb.close()

    if not entries:
        return [_make_row(content="Dữ liệu tiến độ", opinion="Không có dữ liệu WBS.")]

    # Group by 2-level sections: "9.1", "9.2", "4.1", "4" (for top-level)
    sections = {}
    for row_num, wbs, task in entries:
        parts = wbs.split(".")
        if len(parts) >= 2:
            sec = f"{parts[0]}.{parts[1]}"
        else:
            sec = parts[0]
        if sec not in sections:
            sections[sec] = []
        sections[sec].append((row_num, wbs, task))

    # Error lookup by row
    error_by_row = defaultdict(list)
    for e in data_errors:
        if e.row > 0:
            error_by_row[e.row].append(e)

    rows = []
    handled_rows = set()

    for sec in sorted(sections.keys(), key=lambda x: (int(x.split(".")[0]) if x.split(".")[0].isdigit() else 99,
                                                       int(x.split(".")[1]) if "." in x and x.split(".")[1].isdigit() else 0)):
        items = sections[sec]

        # Find section header: the item whose WBS == sec (e.g. "9.1" -> "9.1 HÀNH CHÍNH NHÂN SỰ")
        header_task = ""
        header_wbs = sec
        for _, wbs, task in items:
            if wbs == sec and task:
                header_task = task
                break
        if not header_task:
            for _, wbs, task in items:
                if task:
                    header_task = task
                    break

        # Section header row (BOLD)
        if header_task.startswith(header_wbs + " "):
            header_display = header_task
        else:
            header_display = f"{header_wbs} {header_task}"
        rows.append(_make_row(wbs_text=header_wbs, content=header_display[:80], bold=True))

        # Collect all errors in this section
        section_errors = []  # List of dicts for LLM
        detail_rows = []     # Individual rows (fallback)
        for row_num, wbs, task in items:
            handled_rows.add(row_num)
            item_errs = error_by_row.get(row_num, [])
            if item_errs:
                for e in item_errs:
                    section_errors.append({
                        "wbs": wbs, "task": task[:60] if task else "",
                        "severity": e.severity, "reason": e.reason, "field": e.field,
                    })
                    if task and task.startswith(wbs + " "):
                        ctx = task[:80]
                    elif task:
                        ctx = f"{wbs} {task}"[:80]
                    else:
                        ctx = f"Dòng {row_num} — {e.field}"
                    sl = sev_labels.get(e.severity, "")
                    detail_rows.append(_make_row(content=ctx, opinion=f"{sl} {e.reason}"))

        if not section_errors:
            rows.append(_make_row(content=""))
            continue

        # Use LLM if many errors, else list individually
        if len(section_errors) >= LLM_SUMMARIZE_THRESHOLD:
            try:
                from app.renderer import summarize_section
                llm_text = summarize_section(f"{header_wbs} {header_task}"[:60], section_errors)
                if llm_text:
                    rows.append(_make_row(content="Nhận xét thẩm định", opinion=llm_text))
                    # Add a summary line with counts
                    impact = ", ".join(f"{wbs}" for _, wbs, _ in items[:3] if wbs)
                    rows.append(_make_row(
                        content=f"Các mục bị ảnh hưởng ({len(section_errors)} lỗi)",
                        opinion=f"Bao gồm: {impact}..."
                    ))
                    continue
            except Exception:
                pass

        # Fallback: individual rows
        rows.extend(detail_rows)

    # Orphan errors (rows not in WBS)
    orphans = [e for e in data_errors if e.row > 0 and e.row not in handled_rows]
    if orphans:
        rows.append(_make_row(content="Các dòng không có WBS", bold=True))
        # Group orphans by pattern
        pat_groups = defaultdict(list)
        for e in orphans:
            key = (e.severity, e.reason[:60])
            pat_groups[key].append(e.row)
        for (sev, reason), row_nums in pat_groups.items():
            sl = sev_labels.get(sev, "")
            if len(row_nums) > 5:
                rows.append(_make_row(
                    content=f"{len(row_nums)} dòng",
                    opinion=f"{sl} {reason}\nDòng: {row_nums[0]}-{row_nums[-1]}"
                ))
            else:
                for rn in row_nums:
                    rows.append(_make_row(content=f"Dòng {rn}", opinion=f"{sl} {reason}"))

    return rows


def _fallback_data_rows(data_errors, sev_labels) -> List[List[dict]]:
    """Fallback when Excel can't be opened: list errors by pattern."""
    if not data_errors:
        return [_make_row(content="Dữ liệu tiến độ", opinion="Không có lỗi dữ liệu.")]

    pat_groups = defaultdict(lambda: {"rows": [], "field": ""})
    for e in data_errors:
        key = (e.severity, e.reason[:60])
        pat_groups[key]["rows"].append(e.row)
        pat_groups[key]["field"] = e.field or ""

    rows = []
    for (sev, reason), grp in sorted(pat_groups.items()):
        rl = grp["rows"]
        fd = grp["field"]
        sl = sev_labels.get(sev, "")
        if len(rl) > 5:
            rows.append(_make_row(content=f"{fd} ({len(rl)} dòng)", opinion=f"{sl} {reason}"))
            rows.append(_make_row(content=f"  Dòng: {rl[0]}-{rl[-1]}"))
        else:
            for rn in rl:
                rows.append(_make_row(content=f"Dòng {rn} — {fd}", opinion=f"{sl} {reason}"))
    return rows


def export_word(file_path: str, original_filename: str, reviewer_name: str = "") -> str:
    """Export validation result to Word (.docx) using the BC THẨM ĐỊNH template."""
    result = run_validation(file_path, original_filename)
    errors = result.errors

    cover_errors = [e for e in errors if e.col == "Cover"]
    data_errors = [e for e in errors if e.col != "Cover"]

    project_info = _extract_project_info(original_filename)

    output_path = file_path.rsplit(".", 1)[0] + "_BC_THAM_DINH.docx"
    if os.path.exists(output_path):
        os.remove(output_path)
    shutil.copy2(TEMPLATE_PATH, output_path)

    doc = Document(output_path)

    # Fill header
    header_map = {
        "Tên dự án": project_info.get("ten_du_an", ""),
        "Khu vực": project_info.get("khu_vuc", ""),
        "Nhóm dự án": project_info.get("nhom_du_an", ""),
    }
    rev_match = re.search(r'[Vv]er\s*(\d+)', original_filename)
    if rev_match:
        header_map["Lần thẩm định"] = rev_match.group(1)
    header_map["Ngày hoàn thành thẩm định"] = datetime.now().strftime("%d/%m/%Y")
    file_date_match = re.search(r'(\d{8})', original_filename)
    if file_date_match:
        try:
            dt = datetime.strptime(file_date_match.group(1), "%Y%m%d")
            header_map["Ngày yêu cầu thẩm định"] = dt.strftime("%d/%m/%Y")
        except ValueError:
            pass

    for p in doc.paragraphs:
        text = p.text.strip()
        for key, value in header_map.items():
            if text.startswith(key):
                for r in p.runs:
                    r.text = ""
                if p.runs:
                    p.runs[0].text = f"{key}\t: {value}"
                break

    # Build table
    table = doc.tables[0]
    _remove_rows_from(table, 1)

    # Section A
    _add_row(table, [{"text": "A", "bold": True}, {"text": ""}, {"text": ""},
                      {"text": "PHẦN CHUNG", "bold": True}, {"text": ""}, {"text": ""}, {"text": ""}])
    for row_data in _build_cover_rows(cover_errors):
        _add_row(table, row_data)

    # Section B
    _add_row(table, [{"text": "B", "bold": True}, {"text": ""}, {"text": ""},
                      {"text": "CHI TIẾT", "bold": True}, {"text": ""}, {"text": ""}, {"text": ""}])
    for row_data in _build_wbs_sections(file_path, data_errors):
        _add_row(table, row_data)

    # Sign-off
    if reviewer_name:
        signoff_table = doc.tables[1] if len(doc.tables) > 1 else None
        if signoff_table:
            cell = signoff_table.rows[0].cells[0]
            if cell.paragraphs:
                p = cell.paragraphs[0]
                if "LẬP BÁO CÁO" in p.text.upper():
                    for r in p.runs:
                        r.text = ""
                    if p.runs:
                        p.runs[0].text = f"LẬP BÁO CÁO (GMS)\n{reviewer_name}"

    doc.save(output_path)
    return output_path
